"""PDF to Word converter with layout preservation."""

import asyncio
from pathlib import Path
from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass
from collections import defaultdict

from pdf2image import convert_from_path
from docx import Document
from docx.shared import Pt, Inches, Twips, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image, ImageEnhance, ImageFilter
import pytesseract
import fitz  # PyMuPDF

from app.processors.base import BaseProcessor
from app.core.exceptions import ProcessingError


@dataclass
class TextElement:
    """Represents a text element with position and style."""
    text: str
    x: float
    y: float
    width: float
    height: float
    font_size: float
    font_name: str
    is_bold: bool
    is_italic: bool
    color: Tuple[int, int, int]
    block_no: int
    line_no: int


@dataclass
class LayoutBlock:
    """Represents a block of text with layout info."""
    elements: List[TextElement]
    x: float
    y: float
    width: float
    height: float
    block_type: str  # 'heading', 'paragraph', 'list', 'table'
    column: int  # Column number for multi-column layouts
    indent: float  # Left indentation


class PaddleOCRConverter(BaseProcessor):
    """Convert PDF to Word with layout preservation."""

    QUALITY_PRESETS = {
        "draft": {"dpi": 150, "preprocessing": False},
        "standard": {"dpi": 200, "preprocessing": True},
        "high": {"dpi": 300, "preprocessing": True},
    }

    # Font size to heading level mapping
    HEADING_SIZES = {
        24: 1,  # Heading 1
        20: 1,
        18: 2,  # Heading 2
        16: 2,
        14: 3,  # Heading 3
    }

    def __init__(self, output_dir: Path):
        super().__init__(output_dir)

    async def execute(
        self,
        file: Path,
        output_path: Path,
        lang: str = "eng",
        dpi: int = 200,
        quality: str = "standard",
        preserve_layout: bool = True,
    ) -> Path:
        """
        Convert PDF to Word with layout preservation.

        Args:
            file: Input PDF file path
            output_path: Output DOCX file path
            lang: Language for OCR
            dpi: DPI for image conversion
            quality: Quality preset
            preserve_layout: Preserve document layout

        Returns:
            Path to converted DOCX file
        """
        try:
            # Check if PDF is digital or scanned
            is_digital = await self._is_digital_pdf(file)

            if is_digital and preserve_layout:
                # Use PyMuPDF for digital PDFs (more accurate)
                page_contents = await self._extract_digital_pdf(file)
            else:
                # Use OCR for scanned PDFs
                preset = self.QUALITY_PRESETS.get(quality, self.QUALITY_PRESETS["standard"])
                effective_dpi = dpi if dpi != 200 else preset["dpi"]
                do_preprocessing = preset["preprocessing"]

                images = await self._pdf_to_images(file, effective_dpi)
                page_contents = []

                for i, image in enumerate(images):
                    if do_preprocessing:
                        image = await self._preprocess_image(image)

                    if preserve_layout:
                        content = await self._ocr_with_full_layout(image, i + 1, lang, effective_dpi)
                    else:
                        content = await self._ocr_simple(image, i + 1, lang)

                    page_contents.append(content)

            # Create Word document with layout
            await self._create_docx_with_layout(page_contents, output_path, preserve_layout)

            return output_path

        except ProcessingError:
            raise
        except Exception as e:
            raise ProcessingError(f"OCR conversion failed: {str(e)}")

    async def _is_digital_pdf(self, pdf_path: Path) -> bool:
        """Check if PDF has extractable text (digital) or is scanned."""
        def _check():
            doc = fitz.open(str(pdf_path))
            total_text = 0
            for page in doc:
                text = page.get_text()
                total_text += len(text.strip())
            doc.close()
            # If we have significant text, it's digital
            return total_text > 100

        return await asyncio.to_thread(_check)

    async def _extract_digital_pdf(self, pdf_path: Path) -> List[Dict]:
        """Extract text with layout from digital PDF using PyMuPDF."""
        def _extract():
            doc = fitz.open(str(pdf_path))
            pages = []

            for page_num, page in enumerate(doc):
                page_width = page.rect.width
                page_height = page.rect.height

                # Extract text with detailed info
                blocks_data = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)

                elements = []
                for block in blocks_data.get("blocks", []):
                    if block.get("type") != 0:  # Skip image blocks
                        continue

                    block_no = block.get("number", 0)

                    for line_no, line in enumerate(block.get("lines", [])):
                        for span in line.get("spans", []):
                            text = span.get("text", "").strip()
                            if not text:
                                continue

                            bbox = span.get("bbox", [0, 0, 0, 0])
                            font_size = span.get("size", 11)
                            font_name = span.get("font", "")
                            flags = span.get("flags", 0)
                            color_int = span.get("color", 0)

                            # Parse color
                            r = (color_int >> 16) & 255
                            g = (color_int >> 8) & 255
                            b = color_int & 255

                            elements.append(TextElement(
                                text=text,
                                x=bbox[0],
                                y=bbox[1],
                                width=bbox[2] - bbox[0],
                                height=bbox[3] - bbox[1],
                                font_size=font_size,
                                font_name=font_name,
                                is_bold=bool(flags & 2 ** 4),
                                is_italic=bool(flags & 2 ** 1),
                                color=(r, g, b),
                                block_no=block_no,
                                line_no=line_no,
                            ))

                # Group elements into layout blocks
                layout_blocks = self._group_elements_into_blocks(
                    elements, page_width, page_height
                )

                pages.append({
                    "page": page_num + 1,
                    "width": page_width,
                    "height": page_height,
                    "blocks": layout_blocks,
                    "source": "digital"
                })

            doc.close()
            return pages

        return await asyncio.to_thread(_extract)

    def _group_elements_into_blocks(
        self,
        elements: List[TextElement],
        page_width: float,
        page_height: float,
        is_scanned: bool = False
    ) -> List[LayoutBlock]:
        """Group text elements into logical blocks with layout info."""
        if not elements:
            return []

        # Calculate global font statistics for relative comparison
        all_font_sizes = [e.font_size for e in elements if e.font_size > 0]
        if all_font_sizes:
            # Use median for robustness against outliers
            sorted_sizes = sorted(all_font_sizes)
            median_font_size = sorted_sizes[len(sorted_sizes) // 2]
            avg_font_size_global = sum(all_font_sizes) / len(all_font_sizes)
        else:
            median_font_size = 11
            avg_font_size_global = 11

        # Detect columns by analyzing x positions
        x_positions = [e.x for e in elements]
        columns = self._detect_columns(x_positions, page_width)

        # Group by block number
        blocks_by_num = defaultdict(list)
        for elem in elements:
            blocks_by_num[elem.block_no].append(elem)

        layout_blocks = []
        for block_no, block_elements in sorted(blocks_by_num.items()):
            if not block_elements:
                continue

            # Sort by y then x
            block_elements.sort(key=lambda e: (e.y, e.x))

            # Calculate block bounds
            min_x = min(e.x for e in block_elements)
            min_y = min(e.y for e in block_elements)
            max_x = max(e.x + e.width for e in block_elements)
            max_y = max(e.y + e.height for e in block_elements)

            # Determine column
            col = 0
            for i, (col_start, col_end) in enumerate(columns):
                if min_x >= col_start and min_x < col_end:
                    col = i
                    break

            # Determine block type based on font characteristics
            block_font_sizes = [e.font_size for e in block_elements if e.font_size > 0]
            avg_font_size = sum(block_font_sizes) / len(block_font_sizes) if block_font_sizes else median_font_size
            is_all_bold = all(e.is_bold for e in block_elements)
            first_text = block_elements[0].text if block_elements else ""
            block_text = " ".join(e.text for e in block_elements)

            block_type = "paragraph"

            # For scanned PDFs, use relative font size comparison
            if is_scanned:
                # Heading if font size is 1.3x larger than median AND short text
                is_larger = avg_font_size > median_font_size * 1.3
                is_short = len(block_text) < 150
                is_title_like = (
                    first_text.isupper() or
                    first_text.istitle() or
                    block_text.isupper()
                )
                if is_larger and is_short:
                    block_type = "heading"
                elif is_title_like and is_short and len(block_elements) <= 3:
                    block_type = "heading"
            else:
                # For digital PDFs, use absolute font size
                if avg_font_size >= 14 or (is_all_bold and len(first_text) < 100):
                    block_type = "heading"

            # List detection
            if block_type == "paragraph":
                if first_text.startswith(("•", "-", "●", "○", "■", "▪", "►")):
                    block_type = "list"
                elif len(first_text) >= 2 and first_text[0].isdigit() and first_text[1] in ".):":
                    block_type = "list"

            # Calculate indentation
            indent = min_x - (columns[col][0] if columns else 0)

            layout_blocks.append(LayoutBlock(
                elements=block_elements,
                x=min_x,
                y=min_y,
                width=max_x - min_x,
                height=max_y - min_y,
                block_type=block_type,
                column=col,
                indent=max(0, indent),
            ))

        # Sort blocks by column then y position
        layout_blocks.sort(key=lambda b: (b.column, b.y))

        return layout_blocks

    def _detect_columns(self, x_positions: List[float], page_width: float) -> List[Tuple[float, float]]:
        """Detect column boundaries from x positions."""
        if not x_positions:
            return [(0, page_width)]

        # Simple column detection: check if there's a gap in the middle
        margin = 50  # Typical margin
        content_width = page_width - 2 * margin

        # Check for 2-column layout
        mid_point = page_width / 2
        left_texts = [x for x in x_positions if x < mid_point - 20]
        right_texts = [x for x in x_positions if x > mid_point + 20]

        if left_texts and right_texts:
            left_max = max(left_texts) if left_texts else mid_point
            right_min = min(right_texts) if right_texts else mid_point

            # If there's a significant gap, it's 2 columns
            if right_min - left_max > 30:
                return [
                    (0, mid_point),
                    (mid_point, page_width)
                ]

        return [(0, page_width)]

    async def _pdf_to_images(self, pdf_path: Path, dpi: int) -> List[Image.Image]:
        """Convert PDF pages to PIL Images."""
        def _convert():
            return convert_from_path(str(pdf_path), dpi=dpi)
        return await asyncio.to_thread(_convert)

    async def _preprocess_image(self, image: Image.Image) -> Image.Image:
        """Preprocess image for better OCR accuracy."""
        def _process():
            if image.mode != 'RGB':
                image_rgb = image.convert('RGB')
            else:
                image_rgb = image

            gray = image_rgb.convert('L')

            # Enhance contrast
            enhancer = ImageEnhance.Contrast(gray)
            gray = enhancer.enhance(1.5)

            # Enhance sharpness
            enhancer = ImageEnhance.Sharpness(gray)
            gray = enhancer.enhance(1.2)

            # Denoise
            gray = gray.filter(ImageFilter.MedianFilter(size=1))

            # Binarize
            threshold = 140
            gray = gray.point(lambda x: 255 if x > threshold else 0, mode='1')

            return gray.convert('RGB')

        return await asyncio.to_thread(_process)

    async def _ocr_simple(self, image: Image.Image, page_num: int, lang: str) -> Dict:
        """Simple OCR without layout."""
        def _process():
            config = '--oem 3 --psm 3'
            text = pytesseract.image_to_string(image, lang=lang, config=config)
            lines = [line.strip() for line in text.split('\n') if line.strip()]

            return {
                "page": page_num,
                "blocks": [LayoutBlock(
                    elements=[],
                    x=0, y=0, width=0, height=0,
                    block_type="paragraph",
                    column=0,
                    indent=0,
                )],
                "lines": lines,
                "source": "ocr_simple"
            }

        return await asyncio.to_thread(_process)

    async def _ocr_with_full_layout(
        self,
        image: Image.Image,
        page_num: int,
        lang: str,
        dpi: int
    ) -> Dict:
        """OCR with full layout detection."""
        def _process():
            config = '--oem 3 --psm 3'

            # Get detailed OCR data with positions
            data = pytesseract.image_to_data(
                image, lang=lang, config=config, output_type=pytesseract.Output.DICT
            )

            img_width, img_height = image.size

            # Convert to TextElements
            elements = []
            for i in range(len(data['text'])):
                text = data['text'][i].strip()
                if not text:
                    continue

                conf = int(data['conf'][i])
                if conf < 30:
                    continue

                # Get position (in pixels, need to convert)
                x = data['left'][i]
                y = data['top'][i]
                w = data['width'][i]
                h = data['height'][i]

                # Estimate font size from height (approximate)
                font_size = h * 72 / dpi  # Convert pixels to points

                elements.append(TextElement(
                    text=text,
                    x=x,
                    y=y,
                    width=w,
                    height=h,
                    font_size=font_size,
                    font_name="",
                    is_bold=False,
                    is_italic=False,
                    color=(0, 0, 0),
                    block_no=data['block_num'][i],
                    line_no=data['line_num'][i],
                ))

            # Group into layout blocks (mark as scanned for relative font detection)
            layout_blocks = self._group_elements_into_blocks(elements, img_width, img_height, is_scanned=True)

            return {
                "page": page_num,
                "width": img_width,
                "height": img_height,
                "blocks": layout_blocks,
                "source": "ocr_layout"
            }

        return await asyncio.to_thread(_process)

    async def _create_docx_with_layout(
        self,
        page_contents: List[Dict],
        output_path: Path,
        preserve_layout: bool
    ) -> None:
        """Create Word document preserving layout."""
        def _build_docx():
            doc = Document()

            # Set up styles
            styles = doc.styles

            # Normal style
            normal = styles['Normal']
            normal.font.name = 'Arial'
            normal.font.size = Pt(11)

            # Heading styles
            for i in range(1, 4):
                style_name = f'Heading {i}'
                if style_name in styles:
                    h_style = styles[style_name]
                    h_style.font.name = 'Arial'
                    h_style.font.size = Pt(16 - (i * 2))
                    h_style.font.bold = True

            for page_idx, page_data in enumerate(page_contents):
                source = page_data.get("source", "")

                if source == "ocr_simple":
                    # Simple text output
                    lines = page_data.get("lines", [])
                    for line in lines:
                        if line.strip():
                            doc.add_paragraph(line)
                else:
                    # Layout-aware output
                    blocks = page_data.get("blocks", [])

                    # Handle multi-column by processing column by column
                    columns_blocks = defaultdict(list)
                    for block in blocks:
                        columns_blocks[block.column].append(block)

                    # If multiple columns, use a table for layout
                    num_columns = len(columns_blocks)

                    if num_columns > 1 and preserve_layout:
                        # Create invisible table for columns
                        table = doc.add_table(rows=1, cols=num_columns)
                        table.alignment = WD_TABLE_ALIGNMENT.CENTER

                        for col_idx in range(num_columns):
                            cell = table.cell(0, col_idx)
                            col_blocks = columns_blocks.get(col_idx, [])

                            for block in col_blocks:
                                self._add_block_to_cell(cell, block, preserve_layout)
                    else:
                        # Single column - add directly
                        for block in blocks:
                            self._add_block_to_document(doc, block, preserve_layout)

                # Page break
                if page_idx < len(page_contents) - 1:
                    doc.add_page_break()

            doc.save(str(output_path))

        await asyncio.to_thread(_build_docx)

    def _add_block_to_document(self, doc: Document, block: LayoutBlock, preserve_layout: bool):
        """Add a layout block to the document."""
        if not block.elements and not hasattr(block, 'elements'):
            return

        # Combine elements into lines
        lines = self._elements_to_lines(block.elements)
        if not lines:
            return

        # Determine heading level from font size
        heading_level = 0
        if block.elements:
            avg_size = sum(e.font_size for e in block.elements) / len(block.elements)
            for size, level in sorted(self.HEADING_SIZES.items(), reverse=True):
                if avg_size >= size:
                    heading_level = level
                    break

        for line_text, line_elements in lines:
            if not line_text.strip():
                continue

            p = doc.add_paragraph()

            # Apply style based on block type and font size
            if block.block_type == "heading" or heading_level > 0:
                level = heading_level if heading_level > 0 else 1
                p.style = f'Heading {min(level, 3)}'
            elif block.block_type == "list":
                p.style = 'List Bullet'
            else:
                p.style = 'Normal'

            # Apply indentation if significant
            if preserve_layout and block.indent > 20:
                p.paragraph_format.left_indent = Pt(block.indent / 5)

            # Add text with formatting
            if line_elements and preserve_layout:
                for elem in line_elements:
                    run = p.add_run(elem.text + " ")
                    run.font.size = Pt(min(max(elem.font_size, 8), 24))
                    run.bold = elem.is_bold
                    run.italic = elem.is_italic
                    if elem.color != (0, 0, 0):
                        run.font.color.rgb = RGBColor(*elem.color)
            else:
                p.add_run(line_text)

    def _add_block_to_cell(self, cell, block: LayoutBlock, preserve_layout: bool):
        """Add a layout block to a table cell."""
        lines = self._elements_to_lines(block.elements)
        if not lines:
            return

        for line_text, line_elements in lines:
            if not line_text.strip():
                continue

            p = cell.add_paragraph()

            if block.block_type == "heading":
                p.style = 'Heading 1'
            else:
                p.style = 'Normal'

            if line_elements and preserve_layout:
                for elem in line_elements:
                    run = p.add_run(elem.text + " ")
                    run.font.size = Pt(min(max(elem.font_size, 8), 16))
                    run.bold = elem.is_bold
                    run.italic = elem.is_italic
            else:
                p.add_run(line_text)

    def _elements_to_lines(self, elements: List[TextElement]) -> List[Tuple[str, List[TextElement]]]:
        """Group elements into lines."""
        if not elements:
            return []

        # Group by line number
        lines_dict = defaultdict(list)
        for elem in elements:
            lines_dict[elem.line_no].append(elem)

        lines = []
        for line_no in sorted(lines_dict.keys()):
            line_elements = sorted(lines_dict[line_no], key=lambda e: e.x)
            line_text = " ".join(e.text for e in line_elements)
            lines.append((line_text, line_elements))

        return lines
