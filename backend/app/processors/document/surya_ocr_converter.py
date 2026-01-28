"""PDF to Word converter using Surya OCR for best accuracy."""

import asyncio
from pathlib import Path
from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass, field
from collections import defaultdict
import io

from pdf2image import convert_from_path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import fitz  # PyMuPDF

from app.processors.base import BaseProcessor
from app.core.exceptions import ProcessingError


@dataclass
class TextLine:
    """Represents a detected text line with position."""
    text: str
    x: float
    y: float
    width: float
    height: float
    confidence: float
    polygon: List[Tuple[float, float]] = field(default_factory=list)


@dataclass
class LayoutBlock:
    """Represents a layout block (paragraph, heading, table, etc.)."""
    block_type: str  # "text", "title", "heading", "list", "table", "figure"
    lines: List[TextLine]
    x: float
    y: float
    width: float
    height: float


class SuryaOCRConverter(BaseProcessor):
    """Convert PDF to Word using Surya OCR with layout analysis."""

    QUALITY_PRESETS = {
        "draft": {"dpi": 150},
        "standard": {"dpi": 200},
        "high": {"dpi": 300},
    }

    _foundation_model = None
    _ocr_model = None
    _det_model = None
    _layout_model = None

    def __init__(self, output_dir: Path):
        super().__init__(output_dir)

    @classmethod
    def _load_models(cls):
        """Lazy load Surya models (singleton)."""
        if cls._ocr_model is None:
            from surya.foundation import FoundationPredictor
            from surya.recognition import RecognitionPredictor
            from surya.detection import DetectionPredictor

            cls._foundation_model = FoundationPredictor()
            cls._det_model = DetectionPredictor()
            cls._ocr_model = RecognitionPredictor(cls._foundation_model)

        return cls._det_model, cls._ocr_model

    @classmethod
    def _load_layout_model(cls):
        """Load layout analysis model."""
        if cls._layout_model is None:
            try:
                from surya.layout import LayoutPredictor
                cls._layout_model = LayoutPredictor()
            except (ImportError, Exception):
                cls._layout_model = False  # Mark as unavailable
        return cls._layout_model if cls._layout_model else None

    async def execute(
        self,
        file: Path,
        output_path: Path,
        lang: str = "en",
        quality: str = "standard",
        detect_layout: bool = True,
    ) -> Path:
        """
        Convert PDF to Word using Surya OCR.

        Args:
            file: Input PDF file path
            output_path: Output DOCX file path
            lang: Language code (en, id, zh, ja, ko, etc.)
            quality: Quality preset (draft/standard/high)
            detect_layout: Use layout detection for better structure

        Returns:
            Path to converted DOCX file
        """
        try:
            preset = self.QUALITY_PRESETS.get(quality, self.QUALITY_PRESETS["standard"])
            dpi = preset["dpi"]

            # Check if digital or scanned
            is_digital = await self._is_digital_pdf(file)

            if is_digital:
                # For digital PDFs, use PyMuPDF for text + Surya for layout
                pages = await self._extract_digital_with_surya(file, detect_layout)
            else:
                # For scanned PDFs, use full Surya OCR
                pages = await self._extract_scanned_with_surya(file, lang, dpi, detect_layout)

            # Create Word document
            await self._create_docx(pages, output_path)

            return output_path

        except ProcessingError:
            raise
        except Exception as e:
            raise ProcessingError(f"Surya OCR conversion failed: {str(e)}")

    async def _is_digital_pdf(self, pdf_path: Path) -> bool:
        """Check if PDF has extractable text."""
        def _check():
            doc = fitz.open(str(pdf_path))
            total_text = sum(len(page.get_text().strip()) for page in doc)
            doc.close()
            return total_text > 100
        return await asyncio.to_thread(_check)

    async def _extract_digital_with_surya(self, pdf_path: Path, detect_layout: bool) -> List[Dict]:
        """Extract from digital PDF using PyMuPDF + optional Surya layout."""
        def _extract():
            doc = fitz.open(str(pdf_path))
            pages = []

            for page_num, page in enumerate(doc):
                page_width = page.rect.width
                page_height = page.rect.height

                # Extract text blocks
                blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)

                layout_blocks = []
                for block in blocks.get("blocks", []):
                    if block.get("type") != 0:
                        continue

                    lines = []
                    for line in block.get("lines", []):
                        line_text = ""
                        for span in line.get("spans", []):
                            line_text += span.get("text", "")

                        if line_text.strip():
                            bbox = line.get("bbox", [0, 0, 0, 0])
                            lines.append(TextLine(
                                text=line_text,
                                x=bbox[0],
                                y=bbox[1],
                                width=bbox[2] - bbox[0],
                                height=bbox[3] - bbox[1],
                                confidence=1.0,
                            ))

                    if lines:
                        block_bbox = block.get("bbox", [0, 0, 0, 0])
                        layout_blocks.append(LayoutBlock(
                            block_type="text",
                            lines=lines,
                            x=block_bbox[0],
                            y=block_bbox[1],
                            width=block_bbox[2] - block_bbox[0],
                            height=block_bbox[3] - block_bbox[1],
                        ))

                # Detect block types based on font sizes
                self._detect_block_types(layout_blocks)

                pages.append({
                    "page_num": page_num + 1,
                    "width": page_width,
                    "height": page_height,
                    "blocks": layout_blocks,
                })

            doc.close()
            return pages

        return await asyncio.to_thread(_extract)

    async def _extract_scanned_with_surya(
        self,
        pdf_path: Path,
        lang: str,
        dpi: int,
        detect_layout: bool
    ) -> List[Dict]:
        """Extract from scanned PDF using Surya OCR."""
        def _extract():
            # Convert PDF to images
            pil_images = convert_from_path(str(pdf_path), dpi=dpi)

            # Load models
            det_model, ocr_model = self._load_models()
            layout_model = self._load_layout_model() if detect_layout else None

            # Map language codes
            lang_map = {
                "eng": "en", "ind": "id", "chi_sim": "zh",
                "jpn": "ja", "kor": "ko", "ara": "ar",
                "deu": "de", "fra": "fr", "spa": "es",
            }
            surya_lang = lang_map.get(lang, lang)

            pages = []
            for page_num, pil_image in enumerate(pil_images):
                img_width, img_height = pil_image.size

                # Scale to points (72 dpi)
                scale = 72 / dpi
                page_width = img_width * scale
                page_height = img_height * scale

                # Run OCR with detection (new Surya API)
                ocr_results = ocr_model([pil_image], det_predictor=det_model)

                # Process results
                layout_blocks = []

                if ocr_results and len(ocr_results) > 0:
                    page_result = ocr_results[0]

                    # Group text lines into blocks
                    current_block_lines = []
                    last_y = -100

                    for text_line in page_result.text_lines:
                        bbox = text_line.bbox

                        # Convert to points
                        x = bbox[0] * scale
                        y = bbox[1] * scale
                        w = (bbox[2] - bbox[0]) * scale
                        h = (bbox[3] - bbox[1]) * scale

                        line = TextLine(
                            text=text_line.text,
                            x=x,
                            y=y,
                            width=w,
                            height=h,
                            confidence=text_line.confidence if hasattr(text_line, 'confidence') else 0.9,
                        )

                        # Group into blocks by vertical proximity
                        if abs(y - last_y) > h * 2:
                            if current_block_lines:
                                block = self._lines_to_block(current_block_lines)
                                layout_blocks.append(block)
                            current_block_lines = [line]
                        else:
                            current_block_lines.append(line)

                        last_y = y

                    # Don't forget last block
                    if current_block_lines:
                        block = self._lines_to_block(current_block_lines)
                        layout_blocks.append(block)

                # Run layout analysis if available
                if layout_model and detect_layout:
                    try:
                        layout_results = layout_model([pil_image])
                        if layout_results:
                            self._apply_layout_labels(layout_blocks, layout_results[0], scale)
                    except Exception:
                        pass

                # Fallback: detect types heuristically
                self._detect_block_types(layout_blocks)

                pages.append({
                    "page_num": page_num + 1,
                    "width": page_width,
                    "height": page_height,
                    "blocks": layout_blocks,
                })

            return pages

        return await asyncio.to_thread(_extract)

    def _lines_to_block(self, lines: List[TextLine]) -> LayoutBlock:
        """Convert list of lines to a layout block."""
        if not lines:
            return LayoutBlock("text", [], 0, 0, 0, 0)

        min_x = min(l.x for l in lines)
        min_y = min(l.y for l in lines)
        max_x = max(l.x + l.width for l in lines)
        max_y = max(l.y + l.height for l in lines)

        return LayoutBlock(
            block_type="text",
            lines=lines,
            x=min_x,
            y=min_y,
            width=max_x - min_x,
            height=max_y - min_y,
        )

    def _apply_layout_labels(self, blocks: List[LayoutBlock], layout_result, scale: float):
        """Apply layout labels from Surya layout analysis."""
        if not hasattr(layout_result, 'bboxes'):
            return

        for layout_box in layout_result.bboxes:
            bbox = layout_box.bbox
            label = layout_box.label if hasattr(layout_box, 'label') else 'text'

            lx = bbox[0] * scale
            ly = bbox[1] * scale
            lw = (bbox[2] - bbox[0]) * scale
            lh = (bbox[3] - bbox[1]) * scale

            # Find matching block
            for block in blocks:
                # Check overlap
                if (block.x < lx + lw and block.x + block.width > lx and
                    block.y < ly + lh and block.y + block.height > ly):

                    # Map Surya labels to our types
                    label_map = {
                        "title": "title",
                        "section_header": "heading",
                        "text": "text",
                        "list_item": "list",
                        "table": "table",
                        "figure": "figure",
                        "caption": "caption",
                    }
                    block.block_type = label_map.get(label.lower(), "text")

    def _detect_block_types(self, blocks: List[LayoutBlock]):
        """Heuristically detect block types based on content."""
        if not blocks:
            return

        # Calculate average line height
        all_heights = []
        for block in blocks:
            for line in block.lines:
                all_heights.append(line.height)

        avg_height = sum(all_heights) / len(all_heights) if all_heights else 12

        for block in blocks:
            if block.block_type != "text":
                continue

            if not block.lines:
                continue

            first_line = block.lines[0]
            all_text = " ".join(l.text for l in block.lines)

            # Title: large text, short, at top
            if (first_line.height > avg_height * 1.5 and
                len(all_text) < 100 and
                block.y < 150):
                block.block_type = "title"
                continue

            # Heading: larger text, short
            if (first_line.height > avg_height * 1.3 and
                len(all_text) < 150):
                block.block_type = "heading"
                continue

            # List: starts with bullet or number
            if first_line.text.strip().startswith(("•", "-", "●", "○", "■", "▪", "►", "*")):
                block.block_type = "list"
                continue

            if len(first_line.text) >= 2:
                first_chars = first_line.text[:3].strip()
                if (first_chars[0].isdigit() and
                    len(first_chars) > 1 and
                    first_chars[1] in ".):"):
                    block.block_type = "list"

    async def _create_docx(self, pages: List[Dict], output_path: Path) -> None:
        """Create Word document from extracted content."""
        def _build():
            doc = Document()

            # Set up styles
            styles = doc.styles

            normal = styles['Normal']
            normal.font.name = 'Arial'
            normal.font.size = Pt(11)

            for i in range(1, 4):
                style_name = f'Heading {i}'
                if style_name in styles:
                    h_style = styles[style_name]
                    h_style.font.name = 'Arial'
                    h_style.font.size = Pt(18 - (i * 2))
                    h_style.font.bold = True

            for page_idx, page in enumerate(pages):
                blocks = page.get("blocks", [])

                # Sort blocks by position
                blocks.sort(key=lambda b: (b.y, b.x))

                for block in blocks:
                    if not block.lines:
                        continue

                    # Combine lines into text
                    block_text = "\n".join(l.text for l in block.lines)

                    if not block_text.strip():
                        continue

                    p = doc.add_paragraph()

                    # Apply style based on block type
                    if block.block_type == "title":
                        p.style = 'Title' if 'Title' in styles else 'Heading 1'
                    elif block.block_type == "heading":
                        p.style = 'Heading 1'
                    elif block.block_type == "list":
                        p.style = 'List Bullet' if 'List Bullet' in styles else 'Normal'
                    else:
                        p.style = 'Normal'

                    # Add indentation based on x position
                    if block.x > 50:
                        p.paragraph_format.left_indent = Pt(block.x / 5)

                    # Add text
                    for i, line in enumerate(block.lines):
                        if i > 0:
                            p.add_run('\n')

                        run = p.add_run(line.text)

                        # Estimate font size from line height
                        font_size = max(8, min(line.height * 0.75, 24))
                        run.font.size = Pt(font_size)

                # Page break
                if page_idx < len(pages) - 1:
                    doc.add_page_break()

            doc.save(str(output_path))

        await asyncio.to_thread(_build)
