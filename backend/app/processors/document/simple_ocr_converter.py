"""Simple PDF to Word converter - clean output without fancy formatting."""

import asyncio
from pathlib import Path
from typing import Optional

from app.processors.base import BaseProcessor
from app.core.exceptions import ProcessingError


class SimpleOCRConverter(BaseProcessor):
    """Convert PDF to Word with clean, simple formatting."""

    async def execute(
        self,
        file: Path,
        output_path: Path,
        lang: str = "eng",
        quality: str = "standard",
    ) -> Path:
        """
        Convert PDF to Word with clean output.

        - Single font (Arial 11pt)
        - No bold, italic, or special formatting
        - Clean paragraphs
        - Fast processing
        """
        try:
            import fitz  # PyMuPDF
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from PIL import Image
            import pytesseract
            import io

            def _convert():
                # Quality settings
                dpi_map = {"draft": 150, "standard": 200, "high": 300}
                dpi = dpi_map.get(quality, 200)

                # Open PDF
                pdf = fitz.open(file)
                doc = Document()

                # Set default font for entire document
                style = doc.styles['Normal']
                font = style.font
                font.name = 'Arial'
                font.size = Pt(11)
                font.bold = False
                font.italic = False

                # Set margins
                sections = doc.sections
                for section in sections:
                    section.top_margin = Inches(1)
                    section.bottom_margin = Inches(1)
                    section.left_margin = Inches(1)
                    section.right_margin = Inches(1)

                # Process each page
                for page_num in range(len(pdf)):
                    page = pdf[page_num]

                    # First try to extract text directly (for text-based PDFs)
                    text = page.get_text("text").strip()

                    if not text or len(text) < 50:
                        # Use OCR for scanned pages
                        mat = fitz.Matrix(dpi / 72, dpi / 72)
                        pix = page.get_pixmap(matrix=mat)
                        img_data = pix.tobytes("png")
                        img = Image.open(io.BytesIO(img_data))

                        # OCR
                        text = pytesseract.image_to_string(img, lang=lang)

                    # Clean up text
                    text = self._clean_text(text)

                    # Add text to document
                    if text:
                        # Split into paragraphs
                        paragraphs = text.split('\n\n')
                        for para_text in paragraphs:
                            para_text = para_text.strip()
                            if para_text:
                                # Join lines within paragraph
                                para_text = ' '.join(para_text.split('\n'))
                                para_text = ' '.join(para_text.split())  # Normalize spaces

                                para = doc.add_paragraph()
                                run = para.add_run(para_text)
                                # Ensure clean formatting
                                run.bold = False
                                run.italic = False
                                run.font.name = 'Arial'
                                run.font.size = Pt(11)

                    # Add page break (except for last page)
                    if page_num < len(pdf) - 1:
                        doc.add_page_break()

                pdf.close()

                # Save document
                doc.save(output_path)
                return output_path

            return await asyncio.to_thread(_convert)

        except ImportError as e:
            raise ProcessingError(f"Missing dependency: {e}")
        except Exception as e:
            raise ProcessingError(f"Failed to convert PDF: {str(e)}")

    def _clean_text(self, text: str) -> str:
        """Clean OCR text output."""
        import re

        # Remove multiple spaces
        text = re.sub(r' +', ' ', text)

        # Remove multiple newlines (keep max 2)
        text = re.sub(r'\n{3,}', '\n\n', text)

        # Remove page numbers at bottom (common patterns)
        text = re.sub(r'\n\d+\s*$', '', text)
        text = re.sub(r'^\d+\s*\n', '', text)

        # Remove common OCR artifacts
        text = re.sub(r'[|]{2,}', '', text)
        text = re.sub(r'[-]{5,}', '', text)
        text = re.sub(r'[_]{5,}', '', text)

        return text.strip()
