"""PDF to Word converter using OpenRouter Vision AI (Qwen-VL, etc.)."""

import asyncio
import base64
import io
import re
from pathlib import Path
from typing import Optional, List, Dict

import httpx
from pdf2image import convert_from_path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.processors.base import BaseProcessor
from app.core.exceptions import ProcessingError


class OpenRouterVisionOCR(BaseProcessor):
    """Convert PDF to Word using OpenRouter Vision AI models."""

    # OpenRouter API configuration
    OPENROUTER_BASE_URL = "https://openrouter.ai/api/v1"

    # Available free/cheap models
    AVAILABLE_MODELS = {
        # Free models
        "qwen/qwen2.5-vl-72b-instruct:free": "Qwen2.5-VL 72B (Free)",
        "qwen/qwen-2-vl-7b-instruct:free": "Qwen2-VL 7B (Free)",
        "google/gemini-2.0-flash-exp:free": "Gemini 2.0 Flash (Free)",
        # Cheap models
        "qwen/qwen2.5-vl-72b-instruct": "Qwen2.5-VL 72B",
        "anthropic/claude-3.5-sonnet": "Claude 3.5 Sonnet",
        "openai/gpt-4o-mini": "GPT-4o Mini",
    }

    QUALITY_PRESETS = {
        "draft": {"dpi": 100, "max_tokens": 2048},
        "standard": {"dpi": 150, "max_tokens": 4096},
        "high": {"dpi": 200, "max_tokens": 8192},
    }

    # Document type specific prompts
    DOCUMENT_TYPE_PROMPTS = {
        "general": "",  # Use base prompt only

        "academic": """
ACADEMIC DOCUMENT RULES:
- For two-column layouts: process left column first, then right column
- Preserve citation format exactly: [1], [2], (Author, 2024), etc.
- Keep abstract and keywords sections clearly marked
- Maintain equation formatting as best as possible
- Preserve figure/table captions with their numbers
- Keep footnotes and references section intact""",

        "form": """
FORM DOCUMENT RULES:
- Extract field labels and values as "Label: Value" format
- For checkboxes: use [x] for checked, [ ] for unchecked
- For radio buttons: use (•) for selected, ( ) for unselected
- Preserve form sections and groupings
- Keep signature lines as [SIGNATURE LINE]
- Maintain date fields format""",

        "invoice": """
INVOICE/RECEIPT RULES:
- Extract header: company name, logo description, address, contact
- Capture invoice/receipt number, date, due date
- Format line items as table with: Item, Quantity, Unit Price, Total
- Preserve subtotal, tax, discount, and grand total
- Keep payment terms and bank details
- Note any stamps or signatures as [STAMP] or [SIGNATURE]""",

        "legal": """
LEGAL DOCUMENT RULES:
- Preserve article and section numbering exactly (Article 1, Section 1.1, etc.)
- Keep paragraph numbering and indentation structure
- Maintain defined terms formatting (often in quotes or bold)
- Preserve whereas clauses and recitals
- Keep signature blocks with names and titles
- Note any seals or notary stamps as [SEAL] or [NOTARY STAMP]""",

        "report": """
REPORT DOCUMENT RULES:
- Preserve executive summary section
- Maintain heading hierarchy (Chapter > Section > Subsection)
- Keep bullet points and numbered lists intact
- Preserve charts/graphs as [CHART: description]
- Maintain table of contents structure if present
- Keep page headers and footers info""",

        "letter": """
LETTER/CORRESPONDENCE RULES:
- Preserve letterhead information
- Keep date and reference numbers
- Maintain salutation and closing
- Preserve paragraph structure
- Keep signature block with name and title
- Note any attachments mentioned""",
    }

    def __init__(self, output_dir: Path, api_key: Optional[str] = None):
        super().__init__(output_dir)
        self.api_key = api_key

    async def execute(
        self,
        file: Path,
        output_path: Path,
        model: str = "qwen/qwen2.5-vl-72b-instruct:free",
        quality: str = "standard",
        language: str = "auto",
        document_type: str = "general",
    ) -> Path:
        """
        Convert PDF to Word document using OpenRouter Vision AI.

        Args:
            file: Input PDF file path
            output_path: Output DOCX file path
            model: OpenRouter model to use
            quality: Conversion quality (draft, standard, high)
            language: Output language (auto, en, id, zh, ja, ko, etc.)
            document_type: Type of document for specialized prompts
                - general: Default, works for most documents
                - academic: Research papers, journals, theses
                - form: Forms, applications, questionnaires
                - invoice: Invoices, receipts, bills
                - legal: Contracts, agreements, legal documents
                - report: Business reports, analysis documents
                - letter: Letters, correspondence, memos

        Returns:
            Path to converted DOCX file
        """
        if not self.api_key:
            raise ProcessingError("OpenRouter API key not configured")

        try:
            preset = self.QUALITY_PRESETS.get(quality, self.QUALITY_PRESETS["standard"])
            dpi = preset["dpi"]
            max_tokens = preset["max_tokens"]

            # Convert PDF pages to images
            images = await self._pdf_to_images(file, dpi)

            # Get document type specific prompt
            doc_type_prompt = self.DOCUMENT_TYPE_PROMPTS.get(
                document_type, self.DOCUMENT_TYPE_PROMPTS["general"]
            )

            # Process pages
            page_contents = []
            for i, image in enumerate(images):
                content = await self._extract_content_from_image(
                    image, model, max_tokens, language, doc_type_prompt, i + 1, len(images)
                )
                page_contents.append(content)
                # Add delay between pages to respect rate limits
                if i < len(images) - 1:
                    await asyncio.sleep(0.5)

            # Create Word document from extracted content
            await self._create_docx(page_contents, output_path)

            return output_path

        except ProcessingError:
            raise
        except Exception as e:
            raise ProcessingError(f"OpenRouter Vision OCR failed: {str(e)}")

    async def _pdf_to_images(self, pdf_path: Path, dpi: int) -> List[bytes]:
        """Convert PDF pages to images."""
        def _convert():
            images = convert_from_path(str(pdf_path), dpi=dpi)
            result = []
            for img in images:
                # Resize if too large (max 2000px on longest side)
                max_size = 2000
                if max(img.size) > max_size:
                    ratio = max_size / max(img.size)
                    new_size = (int(img.size[0] * ratio), int(img.size[1] * ratio))
                    img = img.resize(new_size)

                buffer = io.BytesIO()
                img.save(buffer, format="JPEG", quality=85, optimize=True)
                result.append(buffer.getvalue())
            return result

        return await asyncio.to_thread(_convert)

    async def _extract_content_from_image(
        self,
        image_bytes: bytes,
        model: str,
        max_tokens: int,
        language: str,
        doc_type_prompt: str,
        page_num: int,
        total_pages: int
    ) -> Dict:
        """Extract structured content from page image using OpenRouter Vision AI."""

        # Encode image to base64
        base64_image = base64.b64encode(image_bytes).decode("utf-8")

        # Build language instruction
        lang_instruction = ""
        if language != "auto":
            lang_map = {
                "en": "English",
                "id": "Indonesian (Bahasa Indonesia)",
                "zh": "Chinese (中文)",
                "ja": "Japanese (日本語)",
                "ko": "Korean (한국어)",
                "ar": "Arabic (العربية)",
                "de": "German (Deutsch)",
                "fr": "French (Français)",
                "es": "Spanish (Español)",
            }
            lang_name = lang_map.get(language, language)
            lang_instruction = f"\n\nIMPORTANT: Output the content in {lang_name}."

        # Build document type specific section
        doc_type_section = ""
        if doc_type_prompt:
            doc_type_section = f"\n{doc_type_prompt}\n"

        prompt = f"""You are an expert document OCR assistant. Analyze this document page and extract ALL visible text content with precise formatting.

OUTPUT RULES:
1. Extract EVERY word exactly as shown - do not skip any text
2. Use Markdown formatting:
   - # for main titles/headings
   - ## for section headings
   - ### for subheadings
   - **bold** for emphasized text
   - *italic* for italicized text
   - - or * for bullet points
   - 1. 2. 3. for numbered lists
3. For tables, use markdown table format:
   | Column 1 | Column 2 |
   |----------|----------|
   | Data 1   | Data 2   |
4. Mark images/figures as: [IMAGE: brief description]
5. Preserve paragraph structure with blank lines
6. For multi-column layouts, read left-to-right, top-to-bottom
7. Keep original formatting (bold, italic, lists) as shown in the document{doc_type_section}{lang_instruction}

IMPORTANT:
- Do NOT add any commentary, explanation, or notes
- Do NOT translate unless specifically requested
- Output ONLY the document content in markdown format
- Page {page_num} of {total_pages}"""

        try:
            async with httpx.AsyncClient(timeout=120.0) as client:
                response = await client.post(
                    f"{self.OPENROUTER_BASE_URL}/chat/completions",
                    headers={
                        "Authorization": f"Bearer {self.api_key}",
                        "HTTP-Referer": "https://pdfglide.com",
                        "X-Title": "PDFGlide OCR",
                        "Content-Type": "application/json",
                    },
                    json={
                        "model": model,
                        "messages": [
                            {
                                "role": "user",
                                "content": [
                                    {
                                        "type": "image_url",
                                        "image_url": {
                                            "url": f"data:image/jpeg;base64,{base64_image}"
                                        }
                                    },
                                    {
                                        "type": "text",
                                        "text": prompt
                                    }
                                ]
                            }
                        ],
                        "max_tokens": max_tokens,
                        "temperature": 0.1,
                    }
                )

                if response.status_code != 200:
                    error_detail = response.text
                    raise ProcessingError(
                        f"OpenRouter API error (page {page_num}): {response.status_code} - {error_detail}"
                    )

                data = response.json()

                if "error" in data:
                    raise ProcessingError(f"OpenRouter error: {data['error']}")

                content = data["choices"][0]["message"]["content"]
                usage = data.get("usage", {})

                return {
                    "page": page_num,
                    "markdown": content,
                    "tokens_used": usage.get("total_tokens", 0),
                    "model": model,
                }

        except httpx.TimeoutException:
            raise ProcessingError(f"OpenRouter API timeout for page {page_num}")
        except Exception as e:
            if isinstance(e, ProcessingError):
                raise
            raise ProcessingError(f"AI extraction failed for page {page_num}: {str(e)}")

    async def _create_docx(self, page_contents: List[Dict], output_path: Path) -> None:
        """Create Word document from extracted markdown content."""
        def _build_docx():
            doc = Document()

            # Set default font
            style = doc.styles['Normal']
            style.font.name = 'Arial'
            style.font.size = Pt(11)

            for page_data in page_contents:
                markdown = page_data.get("markdown", "")
                self._markdown_to_docx(doc, markdown)

                # Add page break between pages (except last)
                if page_data != page_contents[-1]:
                    doc.add_page_break()

            doc.save(str(output_path))

        await asyncio.to_thread(_build_docx)

    def _markdown_to_docx(self, doc: Document, markdown: str) -> None:
        """Convert markdown content to Word document elements."""
        lines = markdown.split('\n')
        i = 0

        while i < len(lines):
            line = lines[i]

            # Skip empty lines
            if not line.strip():
                i += 1
                continue

            # Headings
            if line.startswith('### '):
                doc.add_heading(line[4:].strip(), level=3)
            elif line.startswith('## '):
                doc.add_heading(line[3:].strip(), level=2)
            elif line.startswith('# '):
                doc.add_heading(line[2:].strip(), level=1)

            # Bullet points
            elif line.strip().startswith('- ') or line.strip().startswith('* '):
                text = line.strip()[2:]
                doc.add_paragraph(
                    self._process_inline_formatting(text),
                    style='List Bullet'
                )

            # Numbered lists
            elif re.match(r'^\d+\.\s', line.strip()):
                text = re.sub(r'^\d+\.\s', '', line.strip())
                doc.add_paragraph(
                    self._process_inline_formatting(text),
                    style='List Number'
                )

            # Tables
            elif line.strip().startswith('|'):
                # Collect all table lines
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith('|'):
                    table_lines.append(lines[i])
                    i += 1
                self._add_table(doc, table_lines)
                continue

            # Image placeholders
            elif line.strip().startswith('[IMAGE:'):
                p = doc.add_paragraph()
                p.add_run(line.strip()).italic = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Regular paragraphs
            else:
                p = doc.add_paragraph()
                self._add_formatted_text(p, line)

            i += 1

    def _process_inline_formatting(self, text: str) -> str:
        """Remove markdown formatting for simple text extraction."""
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = re.sub(r'\*(.+?)\*', r'\1', text)
        text = re.sub(r'__(.+?)__', r'\1', text)
        text = re.sub(r'_(.+?)_', r'\1', text)
        return text

    def _add_formatted_text(self, paragraph, text: str) -> None:
        """Add text with bold/italic formatting to paragraph."""
        pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|[^*]+)'
        parts = re.findall(pattern, text)

        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*'):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            else:
                paragraph.add_run(part)

    def _add_table(self, doc: Document, table_lines: List[str]) -> None:
        """Add a table from markdown table lines."""
        if len(table_lines) < 2:
            return

        # Parse table rows
        rows = []
        for line in table_lines:
            # Skip separator lines (|---|---|)
            if re.match(r'^[\|\s\-:]+$', line):
                continue
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if cells:
                rows.append(cells)

        if not rows:
            return

        # Create table
        num_cols = max(len(row) for row in rows)
        table = doc.add_table(rows=len(rows), cols=num_cols)
        table.style = 'Table Grid'

        for i, row in enumerate(rows):
            for j, cell_text in enumerate(row):
                if j < num_cols:
                    table.rows[i].cells[j].text = self._process_inline_formatting(cell_text)

        # Add spacing after table
        doc.add_paragraph()
