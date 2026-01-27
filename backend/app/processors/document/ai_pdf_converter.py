"""AI-powered PDF to Document converter using Z.ai Vision API."""

import asyncio
import base64
import io
import re
from pathlib import Path
from typing import Optional

from openai import AsyncOpenAI
from pdf2image import convert_from_path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.processors.base import BaseProcessor
from app.core.exceptions import ProcessingError


class AIPDFConverter(BaseProcessor):
    """Convert PDF to Word using Z.ai Vision AI."""

    # Z.ai API configuration
    ZAI_BASE_URL = "https://open.z.ai/api/paas/v4"

    def __init__(self, output_dir: Path, api_key: Optional[str] = None):
        super().__init__(output_dir)
        self.api_key = api_key
        self.client: Optional[AsyncOpenAI] = None

    def _get_client(self) -> AsyncOpenAI:
        """Get or create the OpenAI client configured for Z.ai."""
        if self.client is None:
            if not self.api_key:
                raise ProcessingError("Z.ai API key not configured")
            self.client = AsyncOpenAI(
                api_key=self.api_key,
                base_url=self.ZAI_BASE_URL,
            )
        return self.client

    async def execute(
        self,
        file: Path,
        output_path: Path,
        model: str = "glm-4.6v-flash",
        quality: str = "standard",
        dpi: int = 150,
    ) -> Path:
        """
        Convert PDF to Word document using AI vision.

        Args:
            file: Input PDF file path
            output_path: Output DOCX file path
            model: Z.ai model to use (glm-4.6v-flash, glm-4.6v-flashx, glm-4.6v)
            quality: Conversion quality (draft, standard, high)
            dpi: DPI for PDF to image conversion

        Returns:
            Path to converted DOCX file
        """
        try:
            # Adjust DPI based on quality
            if quality == "draft":
                dpi = 100
            elif quality == "high":
                dpi = 200

            # Convert PDF pages to images
            images = await self._pdf_to_images(file, dpi)

            # Process each page with AI
            page_contents = []
            for i, image in enumerate(images):
                content = await self._extract_content_from_image(image, model, i + 1, len(images))
                page_contents.append(content)

            # Create Word document from extracted content
            await self._create_docx(page_contents, output_path)

            return output_path

        except ProcessingError:
            raise
        except Exception as e:
            raise ProcessingError(f"AI PDF conversion failed: {str(e)}")

    async def _pdf_to_images(self, pdf_path: Path, dpi: int) -> list[bytes]:
        """Convert PDF pages to images."""
        def _convert():
            images = convert_from_path(str(pdf_path), dpi=dpi)
            result = []
            for img in images:
                buffer = io.BytesIO()
                img.save(buffer, format="PNG", optimize=True)
                result.append(buffer.getvalue())
            return result

        return await asyncio.to_thread(_convert)

    async def _extract_content_from_image(
        self,
        image_bytes: bytes,
        model: str,
        page_num: int,
        total_pages: int
    ) -> dict:
        """Extract structured content from page image using AI."""
        client = self._get_client()

        # Encode image to base64
        base64_image = base64.b64encode(image_bytes).decode("utf-8")

        prompt = """Analyze this document page and extract ALL content with precise formatting.

Output as structured markdown following these rules:
1. Preserve ALL text exactly as shown
2. Use # for main headings, ## for subheadings, ### for smaller headings
3. Use **bold** and *italic* where appropriate
4. Use proper bullet points (- or *) and numbered lists (1. 2. 3.)
5. For tables, use markdown table format with | separators
6. Mark images/figures as: [IMAGE: brief description]
7. Preserve paragraph breaks with blank lines
8. For multi-column layouts, process left-to-right, top-to-bottom

Important:
- Extract EVERY word visible on the page
- Maintain the logical reading order
- Keep formatting consistent
- Do NOT add any commentary or explanation, just the content"""

        try:
            response = await client.chat.completions.create(
                model=model,
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/png;base64,{base64_image}"
                                }
                            },
                            {
                                "type": "text",
                                "text": prompt
                            }
                        ]
                    }
                ],
                max_tokens=4096,
                temperature=0.1,
            )

            content = response.choices[0].message.content

            return {
                "page": page_num,
                "markdown": content,
                "tokens_used": response.usage.total_tokens if response.usage else 0
            }

        except Exception as e:
            raise ProcessingError(f"AI extraction failed for page {page_num}: {str(e)}")

    async def _create_docx(self, page_contents: list[dict], output_path: Path) -> None:
        """Create Word document from extracted markdown content."""
        def _build_docx():
            doc = Document()

            # Set default font
            style = doc.styles['Normal']
            style.font.name = 'Arial'
            style.font.size = Pt(11)

            for page_data in page_contents:
                markdown = page_data.get("markdown", "")
                self._markdown_to_docx(doc, markdown)

                # Add page break between pages (except last)
                if page_data != page_contents[-1]:
                    doc.add_page_break()

            doc.save(str(output_path))

        await asyncio.to_thread(_build_docx)

    def _markdown_to_docx(self, doc: Document, markdown: str) -> None:
        """Convert markdown content to Word document elements."""
        lines = markdown.split('\n')
        i = 0

        while i < len(lines):
            line = lines[i]

            # Skip empty lines
            if not line.strip():
                i += 1
                continue

            # Headings
            if line.startswith('### '):
                p = doc.add_heading(line[4:].strip(), level=3)
            elif line.startswith('## '):
                p = doc.add_heading(line[3:].strip(), level=2)
            elif line.startswith('# '):
                p = doc.add_heading(line[2:].strip(), level=1)

            # Bullet points
            elif line.strip().startswith('- ') or line.strip().startswith('* '):
                text = line.strip()[2:]
                p = doc.add_paragraph(self._process_inline_formatting(text), style='List Bullet')

            # Numbered lists
            elif re.match(r'^\d+\.\s', line.strip()):
                text = re.sub(r'^\d+\.\s', '', line.strip())
                p = doc.add_paragraph(self._process_inline_formatting(text), style='List Number')

            # Tables
            elif line.strip().startswith('|'):
                # Collect all table lines
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith('|'):
                    table_lines.append(lines[i])
                    i += 1
                self._add_table(doc, table_lines)
                continue

            # Image placeholders
            elif line.strip().startswith('[IMAGE:'):
                p = doc.add_paragraph()
                p.add_run(line.strip()).italic = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Regular paragraphs
            else:
                p = doc.add_paragraph()
                self._add_formatted_text(p, line)

            i += 1

    def _process_inline_formatting(self, text: str) -> str:
        """Remove markdown formatting for simple text extraction."""
        # Remove bold and italic markers for plain text
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = re.sub(r'\*(.+?)\*', r'\1', text)
        text = re.sub(r'__(.+?)__', r'\1', text)
        text = re.sub(r'_(.+?)_', r'\1', text)
        return text

    def _add_formatted_text(self, paragraph, text: str) -> None:
        """Add text with bold/italic formatting to paragraph."""
        # Pattern to match **bold**, *italic*, or regular text
        pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|[^*]+)'
        parts = re.findall(pattern, text)

        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*'):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            else:
                paragraph.add_run(part)

    def _add_table(self, doc: Document, table_lines: list[str]) -> None:
        """Add a table from markdown table lines."""
        if len(table_lines) < 2:
            return

        # Parse table rows
        rows = []
        for line in table_lines:
            # Skip separator lines (|---|---|)
            if re.match(r'^[\|\s\-:]+$', line):
                continue
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if cells:
                rows.append(cells)

        if not rows:
            return

        # Create table
        num_cols = max(len(row) for row in rows)
        table = doc.add_table(rows=len(rows), cols=num_cols)
        table.style = 'Table Grid'

        for i, row in enumerate(rows):
            for j, cell_text in enumerate(row):
                if j < num_cols:
                    table.rows[i].cells[j].text = self._process_inline_formatting(cell_text)

        # Add spacing after table
        doc.add_paragraph()
