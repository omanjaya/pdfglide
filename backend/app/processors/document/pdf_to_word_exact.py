"""PDF to Word converter with exact layout reproduction."""

import asyncio
import io
import tempfile
from pathlib import Path
from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass, field
from collections import defaultdict

from pdf2image import convert_from_path
from docx import Document
from docx.shared import Pt, Inches, Emu, Twips, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from PIL import Image, ImageEnhance, ImageFilter
import pytesseract
import fitz  # PyMuPDF

from app.processors.base import BaseProcessor
from app.core.exceptions import ProcessingError


@dataclass
class TextBox:
    """Represents a positioned text box."""
    text: str
    x: float  # in points (1/72 inch)
    y: float
    width: float
    height: float
    font_size: float
    font_name: str
    is_bold: bool
    is_italic: bool
    color: Tuple[int, int, int]


@dataclass
class PageContent:
    """Content of a single page."""
    page_num: int
    width: float  # in points
    height: float
    text_boxes: List[TextBox] = field(default_factory=list)
    images: List[Dict] = field(default_factory=list)
    background_image: Optional[bytes] = None  # For scanned pages


class ExactLayoutConverter(BaseProcessor):
    """Convert PDF to Word with exact layout reproduction."""

    QUALITY_PRESETS = {
        "draft": {"dpi": 150},
        "standard": {"dpi": 200},
        "high": {"dpi": 300},
    }

    def __init__(self, output_dir: Path):
        super().__init__(output_dir)

    async def execute(
        self,
        file: Path,
        output_path: Path,
        lang: str = "eng",
        quality: str = "standard",
        include_background: bool = False,
        exact_positions: bool = True,
    ) -> Path:
        """
        Convert PDF to Word with exact layout.

        Args:
            file: Input PDF file path
            output_path: Output DOCX file path
            lang: Language for OCR
            quality: Quality preset (draft/standard/high)
            include_background: Include page image as background (for scanned PDFs)
            exact_positions: Use exact text positioning (text boxes)

        Returns:
            Path to converted DOCX file
        """
        try:
            preset = self.QUALITY_PRESETS.get(quality, self.QUALITY_PRESETS["standard"])
            dpi = preset["dpi"]

            # Check if digital or scanned
            is_digital = await self._is_digital_pdf(file)

            if is_digital:
                pages = await self._extract_digital_pdf(file)
            else:
                pages = await self._extract_scanned_pdf(file, lang, dpi, include_background)

            # Create Word document with exact layout
            await self._create_exact_docx(pages, output_path, exact_positions)

            return output_path

        except Exception as e:
            raise ProcessingError(f"Exact layout conversion failed: {str(e)}")

    async def _is_digital_pdf(self, pdf_path: Path) -> bool:
        """Check if PDF has extractable text."""
        def _check():
            doc = fitz.open(str(pdf_path))
            total_text = sum(len(page.get_text().strip()) for page in doc)
            doc.close()
            return total_text > 100
        return await asyncio.to_thread(_check)

    async def _extract_digital_pdf(self, pdf_path: Path) -> List[PageContent]:
        """Extract content from digital PDF with exact positions."""
        def _extract():
            doc = fitz.open(str(pdf_path))
            pages = []

            for page_num, page in enumerate(doc):
                page_width = page.rect.width
                page_height = page.rect.height

                text_boxes = []

                # Extract text blocks with positions
                blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)

                for block in blocks.get("blocks", []):
                    if block.get("type") != 0:  # Skip image blocks for now
                        continue

                    for line in block.get("lines", []):
                        line_text = ""
                        line_spans = []

                        for span in line.get("spans", []):
                            text = span.get("text", "")
                            if text:
                                line_spans.append(span)
                                line_text += text

                        if line_text.strip() and line_spans:
                            # Use first span's properties for the line
                            first_span = line_spans[0]
                            bbox = line.get("bbox", first_span.get("bbox", [0,0,0,0]))

                            flags = first_span.get("flags", 0)
                            color_int = first_span.get("color", 0)
                            r = (color_int >> 16) & 255
                            g = (color_int >> 8) & 255
                            b = color_int & 255

                            text_boxes.append(TextBox(
                                text=line_text,
                                x=bbox[0],
                                y=bbox[1],
                                width=bbox[2] - bbox[0],
                                height=bbox[3] - bbox[1],
                                font_size=first_span.get("size", 11),
                                font_name=first_span.get("font", "Arial"),
                                is_bold=bool(flags & 2**4),
                                is_italic=bool(flags & 2**1),
                                color=(r, g, b),
                            ))

                # Extract images
                images = []
                for img_index, img in enumerate(page.get_images(full=True)):
                    xref = img[0]
                    try:
                        base_image = doc.extract_image(xref)
                        image_bytes = base_image["image"]

                        # Get image position
                        img_rects = page.get_image_rects(xref)
                        if img_rects:
                            rect = img_rects[0]
                            images.append({
                                "data": image_bytes,
                                "x": rect.x0,
                                "y": rect.y0,
                                "width": rect.width,
                                "height": rect.height,
                                "ext": base_image.get("ext", "png"),
                            })
                    except Exception:
                        pass

                pages.append(PageContent(
                    page_num=page_num + 1,
                    width=page_width,
                    height=page_height,
                    text_boxes=text_boxes,
                    images=images,
                ))

            doc.close()
            return pages

        return await asyncio.to_thread(_extract)

    async def _extract_scanned_pdf(
        self,
        pdf_path: Path,
        lang: str,
        dpi: int,
        include_background: bool
    ) -> List[PageContent]:
        """Extract content from scanned PDF using OCR."""
        def _extract():
            # Convert PDF to images
            pil_images = convert_from_path(str(pdf_path), dpi=dpi)

            pages = []
            for page_num, pil_image in enumerate(pil_images):
                img_width, img_height = pil_image.size

                # Calculate page size in points (72 dpi)
                page_width = img_width * 72 / dpi
                page_height = img_height * 72 / dpi

                # Preprocess for better OCR
                processed = self._preprocess_image(pil_image)

                # Get OCR data with positions
                ocr_data = pytesseract.image_to_data(
                    processed,
                    lang=lang,
                    config='--oem 3 --psm 3',
                    output_type=pytesseract.Output.DICT
                )

                # Group words into lines
                text_boxes = self._ocr_to_textboxes(ocr_data, dpi)

                # Background image if requested
                background = None
                if include_background:
                    img_bytes = io.BytesIO()
                    pil_image.save(img_bytes, format='PNG')
                    background = img_bytes.getvalue()

                pages.append(PageContent(
                    page_num=page_num + 1,
                    width=page_width,
                    height=page_height,
                    text_boxes=text_boxes,
                    background_image=background,
                ))

            return pages

        return await asyncio.to_thread(_extract)

    def _preprocess_image(self, image: Image.Image) -> Image.Image:
        """Preprocess image for OCR."""
        if image.mode != 'RGB':
            image = image.convert('RGB')

        gray = image.convert('L')

        # Enhance contrast
        enhancer = ImageEnhance.Contrast(gray)
        gray = enhancer.enhance(1.5)

        # Sharpen
        enhancer = ImageEnhance.Sharpness(gray)
        gray = enhancer.enhance(1.2)

        return gray.convert('RGB')

    def _ocr_to_textboxes(self, ocr_data: Dict, dpi: int) -> List[TextBox]:
        """Convert OCR data to positioned text boxes."""
        text_boxes = []

        # Group by line
        lines = defaultdict(list)
        for i in range(len(ocr_data['text'])):
            text = ocr_data['text'][i].strip()
            conf = int(ocr_data['conf'][i])

            if text and conf > 30:
                block_num = ocr_data['block_num'][i]
                line_num = ocr_data['line_num'][i]
                key = (block_num, line_num)

                lines[key].append({
                    'text': text,
                    'left': ocr_data['left'][i],
                    'top': ocr_data['top'][i],
                    'width': ocr_data['width'][i],
                    'height': ocr_data['height'][i],
                })

        # Create text box for each line
        scale = 72 / dpi  # Convert pixels to points

        for key in sorted(lines.keys()):
            words = sorted(lines[key], key=lambda w: w['left'])
            if not words:
                continue

            # Combine words into line
            line_text = ' '.join(w['text'] for w in words)

            # Calculate line bounds
            x = words[0]['left'] * scale
            y = min(w['top'] for w in words) * scale
            right = max(w['left'] + w['width'] for w in words) * scale
            bottom = max(w['top'] + w['height'] for w in words) * scale

            # Estimate font size from height
            avg_height = sum(w['height'] for w in words) / len(words)
            font_size = avg_height * scale * 0.75  # Approximate

            text_boxes.append(TextBox(
                text=line_text,
                x=x,
                y=y,
                width=right - x,
                height=bottom - y,
                font_size=max(8, min(font_size, 24)),
                font_name="Arial",
                is_bold=False,
                is_italic=False,
                color=(0, 0, 0),
            ))

        return text_boxes

    async def _create_exact_docx(
        self,
        pages: List[PageContent],
        output_path: Path,
        exact_positions: bool
    ) -> None:
        """Create Word document with exact positioning."""
        def _build():
            doc = Document()

            for page_idx, page in enumerate(pages):
                # Set page size
                section = doc.sections[-1] if page_idx == 0 else doc.add_section()
                section.page_width = Pt(page.width)
                section.page_height = Pt(page.height)
                section.left_margin = Pt(0)
                section.right_margin = Pt(0)
                section.top_margin = Pt(0)
                section.bottom_margin = Pt(0)

                if exact_positions:
                    # Use positioned text boxes
                    self._add_positioned_content(doc, page)
                else:
                    # Use regular paragraphs (simpler but less accurate)
                    self._add_flowing_content(doc, page)

                # Add page break if not last page
                if page_idx < len(pages) - 1:
                    doc.add_page_break()

            doc.save(str(output_path))

        await asyncio.to_thread(_build)

    def _add_positioned_content(self, doc: Document, page: PageContent):
        """Add content with exact positioning using text boxes."""
        # Sort text boxes by position (top to bottom, left to right)
        sorted_boxes = sorted(page.text_boxes, key=lambda b: (b.y, b.x))

        # Group into lines (boxes with similar y position)
        lines = []
        current_line = []
        last_y = -100

        for box in sorted_boxes:
            if abs(box.y - last_y) > 5:  # New line
                if current_line:
                    lines.append(current_line)
                current_line = [box]
                last_y = box.y
            else:
                current_line.append(box)

        if current_line:
            lines.append(current_line)

        # Add each line as a paragraph with appropriate formatting
        for line_boxes in lines:
            line_boxes = sorted(line_boxes, key=lambda b: b.x)

            p = doc.add_paragraph()

            # Calculate indentation based on x position
            if line_boxes:
                first_x = line_boxes[0].x
                if first_x > 10:
                    p.paragraph_format.left_indent = Pt(first_x)

            # Add text with formatting
            for i, box in enumerate(line_boxes):
                # Add space between boxes if needed
                if i > 0:
                    prev_box = line_boxes[i-1]
                    gap = box.x - (prev_box.x + prev_box.width)
                    if gap > 20:  # Significant gap, add tab
                        p.add_run('\t')
                    elif gap > 5:
                        p.add_run(' ')

                run = p.add_run(box.text)
                run.font.size = Pt(box.font_size)
                run.font.bold = box.is_bold
                run.font.italic = box.is_italic

                if box.color != (0, 0, 0):
                    run.font.color.rgb = RGBColor(*box.color)

                # Try to match font
                if box.font_name:
                    font_name = box.font_name.split('-')[0].split('+')[-1]
                    if 'Arial' in font_name or 'Helvetica' in font_name:
                        run.font.name = 'Arial'
                    elif 'Times' in font_name:
                        run.font.name = 'Times New Roman'
                    elif 'Courier' in font_name:
                        run.font.name = 'Courier New'
                    else:
                        run.font.name = 'Arial'

        # Add images
        for img in page.images:
            try:
                # Create a paragraph for the image
                p = doc.add_paragraph()
                if img['x'] > 10:
                    p.paragraph_format.left_indent = Pt(img['x'])

                run = p.add_run()

                # Save image temporarily
                with tempfile.NamedTemporaryFile(suffix=f".{img['ext']}", delete=False) as tmp:
                    tmp.write(img['data'])
                    tmp_path = tmp.name

                # Add image with size
                width_inches = img['width'] / 72
                run.add_picture(tmp_path, width=Inches(width_inches))

                # Clean up
                Path(tmp_path).unlink(missing_ok=True)
            except Exception:
                pass

    def _add_flowing_content(self, doc: Document, page: PageContent):
        """Add content as flowing text (simpler layout)."""
        sorted_boxes = sorted(page.text_boxes, key=lambda b: (b.y, b.x))

        current_text = []
        last_y = -100

        for box in sorted_boxes:
            if abs(box.y - last_y) > 15:  # New paragraph
                if current_text:
                    p = doc.add_paragraph(' '.join(current_text))
                    p.style = 'Normal'
                current_text = [box.text]
                last_y = box.y
            else:
                current_text.append(box.text)

        if current_text:
            doc.add_paragraph(' '.join(current_text))
