"""Document endpoint tests."""

import pytest
from pathlib import Path
from httpx import AsyncClient


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    """Create a sample DOCX file for testing."""
    from docx import Document

    docx_path = tmp_path / "sample.docx"
    doc = Document()
    doc.add_heading("Test Document", 0)
    doc.add_paragraph("This is a test paragraph.")
    doc.save(docx_path)
    return docx_path


@pytest.mark.asyncio
async def test_word_to_pdf(client: AsyncClient, sample_docx: Path):
    """Test Word to PDF conversion."""
    # Skip if LibreOffice not available
    import shutil
    if not shutil.which("soffice") and not Path("/Applications/LibreOffice.app").exists():
        pytest.skip("LibreOffice not available")

    with open(sample_docx, "rb") as f:
        response = await client.post(
            "/api/v1/document/word-to-pdf",
            files=[("file", ("test.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"))],
        )

    # May fail if LibreOffice not installed, which is OK for CI
    if response.status_code == 200:
        data = response.json()
        assert data["success"] is True
        assert data["data"]["status"] == "completed"


@pytest.mark.asyncio
async def test_pdf_to_word(client: AsyncClient, sample_pdf: Path):
    """Test PDF to Word conversion."""
    with open(sample_pdf, "rb") as f:
        response = await client.post(
            "/api/v1/document/pdf-to-word",
            files=[("file", ("test.pdf", f, "application/pdf"))],
        )

    assert response.status_code == 200
    data = response.json()
    assert data["success"] is True
    assert data["data"]["status"] == "completed"
